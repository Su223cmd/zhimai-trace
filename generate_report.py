import os
from docx import Document
from docx.shared import Pt, Cm, Inches, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ============================================================
# 路径定义
# ============================================================
OUTPUT_PATH = r"D:\合同\AI教育比赛\作品材料\更新材料_V2\附2：开发与应用报告.docx"
SCREENSHOT_DIR = r"d:\其它项目\教育（重构版）\docs\screenshots"

SCREENSHOTS = {
    "dashboard": os.path.join(SCREENSHOT_DIR, "dashboard-2026-05-20T04-28-45-194Z.png"),
    "knowledge_center": os.path.join(SCREENSHOT_DIR, "knowledge-center-2026-05-20T04-34-11-017Z.png"),
    "diagnosis_center": os.path.join(SCREENSHOT_DIR, "diagnosis-center-2026-05-20T04-33-29-649Z.png"),
    "agent_monitor": os.path.join(SCREENSHOT_DIR, "agent-monitor-2026-05-20T04-32-50-829Z.png"),
}


# ============================================================
# 辅助函数
# ============================================================
def set_line_spacing(paragraph, spacing_pt=28):
    """设置段落行间距为指定磅值"""
    pPr = paragraph.paragraph_format
    pPr.line_spacing = Pt(spacing_pt)
    pPr.space_before = Pt(0)
    pPr.space_after = Pt(0)


def set_run_font(run, font_name, font_size, bold=False):
    """设置run的字体"""
    run.font.size = Pt(font_size)
    run.bold = bold
    run.font.name = font_name
    # 设置中文字体
    r = run._element
    rPr = r.find(qn('w:rPr'))
    if rPr is None:
        rPr = OxmlElement('w:rPr')
        r.insert(0, rPr)
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), font_name)
    rFonts.set(qn('w:ascii'), font_name)
    rFonts.set(qn('w:hAnsi'), font_name)


def add_title(doc, text):
    """添加文档标题：仿宋三号(16pt)加粗居中"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_line_spacing(p, 28)
    # 首行无缩进
    p.paragraph_format.first_line_indent = Pt(0)
    run = p.add_run(text)
    set_run_font(run, "仿宋_GB2312", 16, bold=True)
    return p


def add_heading1(doc, text):
    """添加一级标题：黑体三号(16pt)"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_line_spacing(p, 28)
    p.paragraph_format.first_line_indent = Pt(0)
    run = p.add_run(text)
    set_run_font(run, "黑体", 16, bold=False)
    return p


def add_heading2(doc, text):
    """添加二级标题：楷体三号(16pt)"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_line_spacing(p, 28)
    p.paragraph_format.first_line_indent = Pt(0)
    run = p.add_run(text)
    set_run_font(run, "楷体_GB2312", 16, bold=False)
    return p


def add_body(doc, text):
    """添加正文段落：仿宋三号(16pt)，首行缩进2字符"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    set_line_spacing(p, 28)
    # 三号字=16pt，2字符缩进=32pt
    p.paragraph_format.first_line_indent = Pt(32)
    run = p.add_run(text)
    set_run_font(run, "仿宋_GB2312", 16, bold=False)
    return p


def add_body_no_indent(doc, text):
    """添加正文段落（无缩进）"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    set_line_spacing(p, 28)
    p.paragraph_format.first_line_indent = Pt(0)
    run = p.add_run(text)
    set_run_font(run, "仿宋_GB2312", 16, bold=False)
    return p


def add_image(doc, image_path, caption_text, width_cm=14):
    """添加截图：居中，带图注"""
    if not os.path.exists(image_path):
        print(f"  [WARN] 图片不存在: {image_path}")
        add_body(doc, f"[图片缺失：{caption_text}]")
        return

    # 插入图片段落（居中）
    p_img = doc.add_paragraph()
    p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_line_spacing(p_img, 28)
    p_img.paragraph_format.first_line_indent = Pt(0)

    run = p_img.add_run()
    run.add_picture(image_path, width=Cm(width_cm))

    # 图注
    p_cap = doc.add_paragraph()
    p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_line_spacing(p_cap, 28)
    p_cap.paragraph_format.first_line_indent = Pt(0)
    run_cap = p_cap.add_run(caption_text)
    set_run_font(run_cap, "仿宋_GB2312", 12, bold=False)  # 图注用小四号(12pt)


def set_page_margins(doc, top=2.54, bottom=2.54, left=3.18, right=3.18):
    """设置页边距（单位：cm）"""
    for section in doc.sections:
        section.top_margin = Cm(top)
        section.bottom_margin = Cm(bottom)
        section.left_margin = Cm(left)
        section.right_margin = Cm(right)


# ============================================================
# 主函数
# ============================================================
def main():
    print("开始生成开发与应用报告...")
    doc = Document()

    # 设置默认页面边距（A4标准）
    set_page_margins(doc)

    # ============ 文档标题 ============
    add_title(doc, "开发与应用报告")

    # ============ 一、开发背景 ============
    add_heading1(doc, "一、开发背景")
    add_body(doc,
        "当前K-12教学中，教师对学生知识状态的判断主要依赖主观经验。"
        "考试后只能知道分数高低，无法精确识别每个学生对具体知识点的掌握程度，"
        "更无法追溯薄弱的根本原因。"
        "CDM认知诊断模型提供了从答题数据客观推断知识状态的理论基础，"
        "但一线教师缺乏将CDM算法落地为可用工具的技术能力。"
        '本项目通过AI辅助编程，开发了一套完整的认知诊断与根因追溯系统——\u201c知脉溯源\u201d。'
    )

    # ============ 二、设计与开发 ============
    add_heading1(doc, "二、设计与开发")

    # （一）平台/技术选择
    add_heading2(doc, "（一）平台/技术选择")
    add_body(doc, "AI平台：Trae（AI编程助手）、VS Code（代码编辑器）")
    add_body(doc, "后端：Python 3.12 + FastAPI + NumPy + SciPy")
    add_body(doc, "前端：React 18 + TypeScript + Vite + Ant Design 5 + @antv/g6")
    add_body(doc, "数据库：PostgreSQL 16（生产）/ SQLite（开发）")
    add_body(doc, "部署：Docker Compose 容器化")
    add_body(doc, "算法：DINA认知诊断模型 + EM期望最大化算法 + 蒙特卡洛EM降级")

    # （二）开发过程
    add_heading2(doc, "（二）开发过程")
    add_body(doc,
        "开发分为8个阶段："
    )
    add_body(doc,
        "Phase 1：基础框架搭建（FastAPI项目结构、数据库ORM模型13张表、DINA模型原型）"
    )
    add_body(doc,
        "Phase 2：CDM核心算法（EM参数估计、蒙特卡洛EM、拟合优度AIC/BIC、反事实推理、遗忘曲线）"
    )
    add_body(doc,
        "Phase 3：Agent框架（5个自治Agent + AgentBus消息总线 + 事件系统，10个发布-订阅事件）"
    )
    add_body(doc,
        "Phase 4：前端页面（13个功能页面 + React Router路由 + Zustand状态管理 + Axios API对接）"
    )
    add_body(doc,
        "Phase 5：组件库（21个可复用组件：MasteryHeatmap、RootCauseTree、KnowledgeGraph等）"
    )
    add_body(doc,
        "Phase 6：测试（后端54用例 + 前端24用例 + TypeScript类型检查）"
    )
    add_body(doc,
        "Phase 7：Docker部署（三服务容器编排 + PostgreSQL + Nginx反向代理）"
    )
    add_body(doc,
        "Phase 8：前端重构（知识中心三栏布局、诊断中心三合一Tab、作业中心聚合、React.lazy懒加载）"
    )

    add_body(doc,
        "在开发过程中，使用Trae AI编程助手辅助完成代码编写、TypeScript类型定义、组件测试等任务。"
        "以下是关键截图："
    )

    # 插入截图
    add_image(doc, SCREENSHOTS["dashboard"], "图1  系统总览Dashboard")
    add_image(doc, SCREENSHOTS["knowledge_center"], "图2  知识中心三栏布局")
    add_image(doc, SCREENSHOTS["diagnosis_center"], "图3  诊断中心三Tab")
    add_image(doc, SCREENSHOTS["agent_monitor"], "图4  Agent监控")

    # （三）功能架构
    add_heading2(doc, "（三）功能架构")
    add_body(doc,
        "知脉溯源围绕"诊断"核心，构建9大功能模块："
    )
    add_body(doc,
        "1. 系统总览 Dashboard —— 项目/课件/诊断统计概览"
    )
    add_body(doc,
        "2. 项目管理 ProjectManage —— 创建项目、配置学科年级"
    )
    add_body(doc,
        "3. 课件管理 CoursewareManage —— PPT课件上传与自动解析"
    )
    add_body(doc,
        "4. 知识中心 KnowledgeCenter —— 三栏布局：左-目录树 + 中-知识图谱 + 右-详情面板"
    )
    add_body(doc,
        "5. 作业中心 HomeworkCenter —— 作业管理 + Q矩阵可视化编辑"
    )
    add_body(doc,
        "6. 诊断中心 DiagnosisCenter —— 三Tab：班级诊断/学生诊断/CDM参数"
    )
    add_body(doc,
        "7. 教学决策 TeachingDecision —— 根因追溯 + 反事实预测 + 学习路径"
    )
    add_body(doc,
        "8. Agent监控 AgentMonitor —— 5个Agent状态实时监控"
    )
    add_body(doc,
        "9. 系统设置 SystemSettings —— 用户配置、CDM阈值、通知偏好"
    )

    add_body(doc,
        "核心技术架构：浏览器 → Nginx反向代理 → FastAPI → 5-Agent协作层 → DINA诊断引擎 → PostgreSQL"
    )

    # ============ 三、应用过程与效果 ============
    add_heading1(doc, "三、应用过程与效果")

    add_body(doc,
        "应用场景：教师创建诊断项目（如"高一地理诊断"）"
        "→ 上传PPT课件自动解析知识点 → 创建作业录入题目和Q矩阵 "
        "→ 导入学生答题数据 → 系统执行CDM诊断 → 查看诊断结果和根因追溯。"
    )
    add_body(doc, "教学效益：")
    add_body(doc,
        "1）诊断效率：从人工分析需要2-3小时缩短为数秒自动完成"
    )
    add_body(doc,
        "2）定位精度：从"班级整体不行"的模糊判断，"
        "精确到"XX学生在XX知识点上掌握率XX%""
    )
    add_body(doc,
        "3）根因追溯：自动沿知识图谱前驱路径追溯深层原因，替代教师凭经验推测"
    )
    add_body(doc,
        "4）数据驱动决策：反事实预测量化补课收益，个性化学习路径减少无效复习"
    )

    # ============ 四、创新与反思 ============
    add_heading1(doc, "四、创新与反思")

    add_body(doc, "创新点：")
    add_body(doc,
        "1）CDM+知识图谱联动：将DINA诊断结果与6层知识图谱的前驱关系结合，实现深层根因追溯"
    )
    add_body(doc,
        "2）5-Agent协作：诊断/知识/追溯/教学/进化Agent通过AgentBus协同，"
        "自动化诊断→决策链路"
    )
    add_body(doc,
        "3）反事实推理：基于DINA参数化模型，"
        "量化"如果补上X，Y将从A%→B%"的预测"
    )

    add_body(doc,
        "改进方向：当前仅支持DINA单模型，未来可扩展DINO/HO-DINA多模型对比；"
        "课标数据仍需手动导入，计划接入更多课程标准的自动解析。"
    )

    # 保存文档
    doc.save(OUTPUT_PATH)
    print(f"报告已成功生成：{OUTPUT_PATH}")


if __name__ == "__main__":
    main()